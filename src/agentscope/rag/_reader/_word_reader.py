# -*- coding: utf-8 -*-
"""The Word reader to read and chunk Word documents."""
import base64
import hashlib
from typing import Any, Literal

from ._reader_base import ReaderBase
from ._text_reader import TextReader
from .._document import Document, DocMetadata
from ...message import ImageBlock, Base64Source


class WordReader(ReaderBase):
    """The Word reader that splits text into chunks by a fixed chunk size."""

    def __init__(
        self,
        chunk_size: int = 512,
        split_by: Literal["char", "sentence", "paragraph"] = "sentence",
        include_image: bool = False,
        separate_table: bool = False,
    ) -> None:
        """Initialize the Word reader.

        Args:
            chunk_size (`int`, default to 512):
                The size of each chunk, in number of characters.
            split_by (`Literal["char", "sentence", "paragraph"]`, default to \
            "sentence"):
                The unit to split the text, can be "char", "sentence", or
                "paragraph". The "sentence" option is implemented using the
                "nltk" library, which only supports English text.
            include_image (`bool`, default to False):
                Whether to include image content in the document. If True,
                images will be extracted and included as text descriptions.
            separate_table (`bool`, default to False):
                Whether to treat tables as separate documents. If True,
                each table will be extracted as a separate Document object
                instead of being processed as regular text.
        """
        if chunk_size <= 0:
            raise ValueError(
                f"The chunk_size must be positive, got {chunk_size}",
            )

        if split_by not in ["char", "sentence", "paragraph"]:
            raise ValueError(
                "The split_by must be one of 'char', 'sentence' or "
                f"'paragraph', got {split_by}",
            )

        self.chunk_size = chunk_size
        self.split_by = split_by
        self.include_image = include_image
        self.separate_table = separate_table

        # To avoid code duplication, we use TextReader to do the chunking.
        self._text_reader = TextReader(
            self.chunk_size,
            self.split_by,
        )

    async def __call__(
        self,
        word_path: str,
    ) -> list[Document]:
        """Read a Word document, split it into chunks, and return a list of
        Document objects.

        Args:
            word_path (`str`):
                The input Word document file path (.docx file).

        Returns:
            `list[Document]`:
                A list of Document objects, where the metadata contains the
                chunked text, doc id and chunk id.
        """
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise ImportError(
                "Please install python-docx to use the Word reader. "
                "You can install it by `pip install python-docx`.",
            ) from e

        # Load the Word document
        doc = DocxDocument(word_path)

        # Extract content in order (paragraphs, tables, and images)
        # We'll collect all content pieces in order, then process them
        content_pieces = []  # List of (type, content) tuples: ('text',
        # text) or ('image', image_doc)

        # Process all elements in document order using a more reliable approach
        # Create mappings for efficient lookup
        paragraph_elements = {p._element: p for p in doc.paragraphs}
        table_elements = {t._element: t for t in doc.tables}

        # Create a set of all drawing elements that are part of paragraphs
        processed_drawing_elements = set()
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Check if this run contains drawing elements
                for elem in run._element.iter():
                    if elem.tag.split("}")[-1] == "drawing":
                        processed_drawing_elements.add(elem)

        # Process elements in document order
        for element in doc.element.body:
            tag_name = self._get_tag_name(element)
            self._process_element_by_type_ordered(
                element,
                tag_name,
                paragraph_elements,
                table_elements,
                processed_drawing_elements,
                content_pieces,
                word_path,
            )

        # Generate document ID
        doc_id = self.get_doc_id(word_path)

        # Process content pieces in order
        return await self._process_content_pieces(content_pieces, doc_id)

    def _get_tag_name(self, element: Any) -> str:
        """Get tag name from element.

        Args:
            element (`Any`):
                XML element object.

        Returns:
            `str`:
                The tag name.
        """
        return (
            element.tag.split("}")[-1] if "}" in element.tag else element.tag
        )

    def _process_element_by_type_ordered(
        self,
        element: Any,
        tag_name: str,
        paragraph_elements: dict,
        table_elements: dict,
        processed_drawing_elements: set,
        content_pieces: list,
        word_path: str,
    ) -> None:
        """Process element based on its type and maintain order.

        Args:
            element (`Any`):
                XML element from Word document.
            tag_name (`str`):
                The tag name of the element.
            paragraph_elements (`dict`):
                Dictionary mapping paragraph elements to paragraph objects.
            table_elements (`dict`):
                Dictionary mapping table elements to table objects.
            processed_drawing_elements (`set`):
                Set of drawing elements that have been processed.
            content_pieces (`list`):
                List to collect content pieces.
            word_path (`str`):
                The path to the Word document.
        """
        if tag_name == "p":
            self._process_paragraph_element_ordered(
                element,
                paragraph_elements,
                content_pieces,
                word_path,
            )
        elif tag_name == "tbl":
            self._process_table_element_ordered(
                element,
                table_elements,
                content_pieces,
                word_path,
            )
        elif tag_name == "drawing" and self.include_image:
            self._process_drawing_element_ordered(
                element,
                processed_drawing_elements,
                content_pieces,
                word_path,
            )

    def _process_element_by_type(
        self,
        element: Any,
        tag_name: str,
        paragraph_elements: dict,
        table_elements: dict,
        processed_drawing_elements: set,
        text_content_parts: list,
        image_documents: list,
        word_path: str,
    ) -> None:
        """Process element based on its type.

        Args:
            element (`Any`):
                XML element from Word document.
            tag_name (`str`):
                The tag name of the element.
            paragraph_elements (`dict`):
                Dictionary mapping paragraph elements to paragraph objects.
            table_elements (`dict`):
                Dictionary mapping table elements to table objects.
            processed_drawing_elements (`set`):
                Set of drawing elements that have been processed.
            text_content_parts (`list`):
                List to collect text content parts.
            image_documents (`list`):
                List to collect image documents.
            word_path (`str`):
                The path to the Word document.
        """
        if tag_name == "p":
            self._process_paragraph_element(
                element,
                paragraph_elements,
                text_content_parts,
                image_documents,
                word_path,
            )
        elif tag_name == "tbl":
            self._process_table_element(
                element,
                table_elements,
                text_content_parts,
            )
        elif tag_name == "drawing" and self.include_image:
            self._process_drawing_element(
                element,
                processed_drawing_elements,
                image_documents,
                word_path,
            )

    def _process_paragraph_element(
        self,
        element: Any,
        paragraph_elements: dict,
        text_content_parts: list,
        image_documents: list,
        word_path: str,
    ) -> None:
        """Process paragraph element.

        Args:
            element (`Any`):
                XML element representing a paragraph.
            paragraph_elements (`dict`):
                Dictionary mapping paragraph elements to paragraph objects.
            text_content_parts (`list`):
                List to collect text content parts.
            image_documents (`list`):
                List to collect image documents.
            word_path (`str`):
                The path to the Word document.
        """
        paragraph = paragraph_elements.get(element)
        if not paragraph:
            return

        text = paragraph.text.strip()
        if text:
            text_content_parts.append(text)

        # Always process images if include_image is True, even if paragraph
        # text is empty
        if self.include_image:
            self._process_paragraph_images(
                paragraph,
                image_documents,
                word_path,
            )

    def _process_paragraph_images(
        self,
        paragraph: Any,
        image_documents: list,
        word_path: str,
    ) -> None:
        """Process images in paragraph.

        Args:
            paragraph (`Any`):
                Paragraph object from python-docx.
            image_documents (`list`):
                List to collect image documents.
            word_path (`str`):
                The path to the Word document.
        """
        for run in paragraph.runs:
            # pylint: disable=protected-access
            image_doc = self._extract_inline_image_document(
                run._element,
                word_path,
            )
            if image_doc:
                image_documents.append(image_doc)

    def _process_table_element(
        self,
        element: Any,
        table_elements: dict,
        content_parts: list,
    ) -> None:
        """Process table element.

        Args:
            element (`Any`):
                XML element representing a table.
            table_elements (`dict`):
                Dictionary mapping table elements to table objects.
            content_parts (`list`):
                List to collect content parts.
        """
        table = table_elements.get(element)
        if not table:
            return

        table_texts = []
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                table_texts.append(" | ".join(row_texts))

        if table_texts:
            table_content = "TABLE:\n" + "\n".join(table_texts)
            content_parts.append(table_content)

    def _process_drawing_element(
        self,
        element: Any,
        processed_drawing_elements: set,
        image_documents: list,
        word_path: str,
    ) -> None:
        """Process standalone drawing element.

        Args:
            element (`Any`):
                XML drawing element.
            processed_drawing_elements (`set`):
                Set of drawing elements that have been processed.
            image_documents (`list`):
                List to collect image documents.
            word_path (`str`):
                The path to the Word document.
        """
        if element not in processed_drawing_elements:
            image_doc = self._extract_drawing_image_document(
                element,
                word_path,
            )
            if image_doc:
                image_documents.append(image_doc)

    def _extract_inline_image_document(
        self,
        run_element: Any,
        word_path: str,
    ) -> Document | None:
        """Extract image document from inline image in a run element.

        Args:
            run_element: The run element that may contain images.
            word_path: The path to the Word document.

        Returns:
            A Document object containing the image, or None if no image found.
        """
        try:
            # Look for image references in the run element
            for elem in run_element.iter():
                if elem.tag.endswith("blip"):
                    # Found an image reference
                    r_embed = elem.get(
                        "{http://schemas.openxmlformats.org/officeDocument/"
                        "2006/relationships}embed",
                    )
                    if r_embed:
                        return self._create_image_document_from_relationship(
                            r_embed,
                            word_path,
                        )
        except Exception:
            pass
        return None

    def _extract_drawing_image_document(
        self,
        drawing_element: Any,
        word_path: str,
    ) -> Document | None:
        """Extract image document from a drawing element.

        Args:
            drawing_element (`Any`):
                The drawing element from the Word document.
            word_path (`str`):
                The path to the Word document.

        Returns:
            `Document | None`:
                A Document object containing the image,
                 or None if no image found.
        """
        try:
            # Look for image references in the drawing element
            for elem in drawing_element.iter():
                if elem.tag.endswith("blip"):
                    # Found an image reference
                    r_embed = elem.get(
                        "{http://schemas.openxmlformats.org/officeDocument/"
                        "2006/relationships}embed",
                    )
                    if r_embed:
                        return self._create_image_document_from_relationship(
                            r_embed,
                            word_path,
                        )
        except Exception:
            pass
        return None

    def _create_image_document_from_relationship(
        self,
        relationship_id: str,
        word_path: str,
    ) -> Document | None:
        """Create an image document from a relationship ID.

        Args:
            relationship_id (`str`):
                The relationship ID of the image.
            word_path (`str`):
                The path to the Word document.

        Returns:
            `Document | None`:
                A Document object containing the image, or None if extraction
                fails.
        """
        try:
            from docx import Document as DocxDocument
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            # Load the Word document to access relationships
            doc = DocxDocument(word_path)

            # Find the relationship
            for rel in doc.part.rels.values():
                if rel.rId == relationship_id and rel.reltype == RT.IMAGE:
                    # Get the image part
                    image_part = rel.target_part
                    if image_part:
                        # Get the image data
                        image_data = image_part.blob
                        if image_data:
                            # Determine media type based on content type
                            content_type = image_part.content_type
                            media_type = (
                                self._get_media_type_from_content_type(
                                    content_type,
                                )
                            )

                            # Convert to base64
                            base64_data = base64.b64encode(image_data).decode(
                                "utf-8",
                            )

                            # Create ImageBlock
                            image_block = ImageBlock(
                                type="image",
                                source=Base64Source(
                                    type="base64",
                                    media_type=media_type,
                                    data=base64_data,
                                ),
                            )

                            # Create Document
                            doc_id = self.get_doc_id(word_path)
                            return Document(
                                metadata=DocMetadata(
                                    content=image_block,
                                    doc_id=doc_id,
                                    chunk_id=0,
                                    total_chunks=1,
                                ),
                            )
        except Exception:
            # If extraction fails, return None
            pass
        return None

    def _get_media_type_from_content_type(self, content_type: str) -> str:
        """Get media type from content type.

        Args:
            content_type (`str`):
                The content type from the image part.

        Returns:
            `str`:
                The media type string.
        """
        # Map common content types to media types
        content_type_mapping = {
            "image/jpeg": "image/jpeg",
            "image/png": "image/png",
            "image/gif": "image/gif",
            "image/bmp": "image/bmp",
            "image/tiff": "image/tiff",
            "image/webp": "image/webp",
        }

        return content_type_mapping.get(content_type, "image/jpeg")

    def get_doc_id(self, word_path: str) -> str:
        """Get the document ID.

        This function can be used to check if the doc_id already exists in the
        knowledge base.

        Args:
            word_path (`str`):
                The path to the Word document file.

        Returns:
            `str`:
                A unique document ID for the Word document.
        """
        return hashlib.sha256(word_path.encode("utf-8")).hexdigest()

    def _process_paragraph_element_ordered(
        self,
        element: Any,
        paragraph_elements: dict,
        content_pieces: list,
        word_path: str,
    ) -> None:
        """Process paragraph element and maintain order.

        Args:
            element (`Any`):
                XML element representing a paragraph.
            paragraph_elements (`dict`):
                Dictionary mapping paragraph elements to paragraph objects.
            content_pieces (`list`):
                List to collect content pieces.
            word_path (`str`):
                The path to the Word document.
        """
        paragraph = paragraph_elements.get(element)
        if not paragraph:
            return

        text = paragraph.text.strip()
        if text:
            content_pieces.append(("text", text))

        # Always process images if include_image is True, even if paragraph
        # text is empty
        if self.include_image:
            image_doc = self._process_paragraph_images_ordered(
                paragraph,
                word_path,
            )
            if image_doc:
                content_pieces.append(("image", image_doc))

    def _process_table_element_ordered(
        self,
        element: Any,
        table_elements: dict,
        content_pieces: list,
        word_path: str,
    ) -> None:
        """Process table element and maintain order.

        Args:
            element (`Any`):
                XML element representing a table.
            table_elements (`dict`):
                Dictionary mapping table elements to table objects.
            content_pieces (`list`):
                List to collect content pieces.
            word_path (`str`):
                The path to the Word document.
        """
        table = table_elements.get(element)
        if not table:
            return

        if self.separate_table:
            # Create a separate table document
            table_doc = self._create_table_document(table, word_path)
            if table_doc:
                content_pieces.append(("table", table_doc))
        else:
            # Extract table content as text
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))

            if table_text:
                content_pieces.append(("text", "\n".join(table_text)))

    def _process_drawing_element_ordered(
        self,
        element: Any,
        processed_drawing_elements: set,
        content_pieces: list,
        word_path: str,
    ) -> None:
        """Process drawing element and maintain order.

        Args:
            element (`Any`):
                XML drawing element.
            processed_drawing_elements (`set`):
                Set of drawing elements that have been processed.
            content_pieces (`list`):
                List to collect content pieces.
            word_path (`str`):
                The path to the Word document.
        """
        if element not in processed_drawing_elements:
            image_doc = self._extract_drawing_image_document(
                element,
                word_path,
            )
            if image_doc:
                content_pieces.append(("image", image_doc))

    def _process_paragraph_images_ordered(
        self,
        paragraph: Any,
        word_path: str,
    ) -> Document | None:
        """Process images in paragraph and maintain order.

        Args:
            paragraph (`Any`):
                Paragraph object from python-docx.
            word_path (`str`):
                The path to the Word document.

        Returns:
            `Document | None`:
                The first image document found, or None if no images.
        """
        for run in paragraph.runs:
            # pylint: disable=protected-access
            image_doc = self._extract_inline_image_document(
                run._element,
                word_path,
            )
            if image_doc:
                return image_doc
        return None

    def _create_table_document(
        self,
        table: Any,
        word_path: str,
    ) -> Document | None:
        """Create a separate document for a table.

        Args:
            table (`Any`):
                The table object from python-docx.
            word_path (`str`):
                The path to the Word document.

        Returns:
            `Document | None`:
                A Document object containing the table data, or None if table
                is empty.
        """
        try:
            # Extract table content
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)

            if not table_data or not any(
                any(cell for cell in row) for row in table_data
            ):
                return None

            # Create table text representation
            table_text = []
            for row in table_data:
                if any(cell for cell in row):  # Skip empty rows
                    table_text.append(" | ".join(cell for cell in row if cell))

            if not table_text:
                return None

            # Create TextBlock for the table
            from ...message import TextBlock

            table_block = TextBlock(
                type="text",
                text="\n".join(table_text),
            )

            # Create Document
            doc_id = self.get_doc_id(word_path)
            return Document(
                metadata=DocMetadata(
                    content=table_block,
                    doc_id=doc_id,
                    chunk_id=0,  # Tables are treated as single chunks
                    total_chunks=1,
                ),
            )

        except Exception:
            # If table processing fails, return None
            pass
        return None

    async def _process_content_pieces(
        self,
        content_pieces: list,
        doc_id: str,
    ) -> list[Document]:
        """Process content pieces in order and return documents.

        Args:
            content_pieces (`list`):
                List of (type, content) tuples from document elements.
            doc_id (`str`):
                The document ID.

        Returns:
            `list[Document]`:
                A list of Document objects processed from content pieces.
        """
        all_docs = []
        current_text_parts = []

        for piece_type, content in content_pieces:
            if piece_type == "text":
                current_text_parts.append(content)
            elif piece_type == "image":
                await self._process_accumulated_text(
                    current_text_parts,
                    all_docs,
                    doc_id,
                )
                all_docs.append(content)
            elif piece_type == "table":
                await self._process_accumulated_text(
                    current_text_parts,
                    all_docs,
                    doc_id,
                )
                all_docs.append(content)

        # Process any remaining text
        await self._process_accumulated_text(
            current_text_parts,
            all_docs,
            doc_id,
        )

        return all_docs

    async def _process_accumulated_text(
        self,
        current_text_parts: list,
        all_docs: list,
        doc_id: str,
    ) -> None:
        """Process accumulated text parts and add to all_docs.

        Args:
            current_text_parts (`list`):
                List of text parts that have been accumulated.
            all_docs (`list`):
                List to collect all documents.
            doc_id (`str`):
                The document ID.
        """
        if current_text_parts:
            full_text = "\n\n".join(current_text_parts)
            if full_text.strip():
                text_docs = await self._text_reader(full_text)
                for doc_obj in text_docs:
                    doc_obj.id = doc_id
                all_docs.extend(text_docs)
            current_text_parts.clear()
